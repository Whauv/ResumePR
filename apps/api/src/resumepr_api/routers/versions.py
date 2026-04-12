from __future__ import annotations

import io
import json
from pathlib import Path

from docx import Document
from fastapi import APIRouter, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from resumepr_api.models.schemas import (
    ParsedResume,
    ResumeVersion,
    ResumeVersionMetadata,
    VersionDiffItem,
    VersionDiffResponse,
    VersionListResponse,
    VersionSummary,
)
from resumepr_api.services.db import get_connection

router = APIRouter(prefix="/api/versions", tags=["versions"])


def row_to_version(row) -> ResumeVersion:
    return ResumeVersion(
        version_id=row[0],
        base_resume_id=row[1],
        metadata=ResumeVersionMetadata(
            version_number=row[3],
            job_id=row[4],
            company_name=row[5] or "",
            role=row[6] or "",
            accepted_count=row[7],
            rejected_count=row[8],
            timestamp=row[10],
            ats_score_before=float(row[11] or 0),
            ats_score_after=float(row[12] or 0),
        ),
        resume_json=ParsedResume.model_validate(json.loads(row[9])),
    )


def flatten_resume_for_diff(resume: ParsedResume) -> dict[str, str]:
    flattened = {
        "summary.text": resume.summary,
        "skills.list": ", ".join(resume.skills),
    }
    for index, item in enumerate(resume.experience):
        prefix = f"experience[{index}]"
        flattened[f"{prefix}.header"] = " - ".join([value for value in [item.title, item.company, item.dates] if value])
        for bullet_index, bullet in enumerate(item.bullets):
            flattened[f"{prefix}.bullet[{bullet_index}]"] = bullet
    for index, item in enumerate(resume.education):
        flattened[f"education[{index}]"] = " - ".join([value for value in [item.degree, item.institution, item.dates] if value])
    return flattened


def diff_versions(current: ResumeVersion, previous: ResumeVersion | None) -> VersionDiffResponse:
    current_map = flatten_resume_for_diff(current.resume_json)
    previous_map = flatten_resume_for_diff(previous.resume_json) if previous else {}
    keys = sorted(set(current_map) | set(previous_map))
    changes = []
    for key in keys:
        before = previous_map.get(key, "")
        after = current_map.get(key, "")
        if before == after:
            continue
        section, field = key.split(".", 1) if "." in key else (key, "text")
        changes.append(VersionDiffItem(section=section, field=field, before=before, after=after))

    return VersionDiffResponse(
        version_id=current.version_id,
        previous_version_id=previous.version_id if previous else None,
        changes=changes,
    )


def draw_pdf_resume(buffer: io.BytesIO, resume: ParsedResume, template: str) -> None:
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    margin = 54
    y = height - margin

    accent = {
        "modern": (0.01, 0.41, 0.44),
        "classic": (0.12, 0.12, 0.12),
        "minimal": (0.35, 0.35, 0.35),
    }.get(template, (0.01, 0.41, 0.44))

    pdf.setFillColorRGB(*accent)
    pdf.setFont("Helvetica-Bold", 20)
    pdf.drawString(margin, y, resume.name or "Resume")
    y -= 24

    pdf.setFillColorRGB(0.2, 0.2, 0.2)
    pdf.setFont("Helvetica", 10)
    contact = " | ".join(f"{key}: {value}" for key, value in resume.contact.items())
    for line in [contact, "", "Summary", resume.summary, "", "Skills", ", ".join(resume.skills)]:
        if not line:
            y -= 10
            continue
        for chunk in [line[i:i + 95] for i in range(0, len(line), 95)]:
            pdf.drawString(margin, y, chunk)
            y -= 14
            if y < margin:
                pdf.showPage()
                y = height - margin
                pdf.setFont("Helvetica", 10)

    for item in resume.experience:
        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(margin, y, " - ".join([value for value in [item.title, item.company, item.dates] if value]))
        y -= 14
        pdf.setFont("Helvetica", 10)
        for bullet in item.bullets:
            for chunk in [bullet[i:i + 90] for i in range(0, len(bullet), 90)]:
                pdf.drawString(margin + 10, y, f"* {chunk}" if chunk == bullet[:90] else f"  {chunk}")
                y -= 14
                if y < margin:
                    pdf.showPage()
                    y = height - margin
                    pdf.setFont("Helvetica", 10)

    pdf.save()


def build_docx_resume(resume: ParsedResume, template: str) -> io.BytesIO:
    document = Document()
    title = document.add_heading(resume.name or "Resume", level=0)
    if template == "minimal":
        title.style.font.name = "Arial"

    if resume.contact:
        document.add_paragraph(" | ".join(f"{key}: {value}" for key, value in resume.contact.items()))

    if resume.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(resume.summary)

    if resume.skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(resume.skills))

    if resume.experience:
        document.add_heading("Experience", level=1)
        for item in resume.experience:
            document.add_paragraph(" - ".join([value for value in [item.title, item.company, item.dates] if value]))
            for bullet in item.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    if resume.education:
        document.add_heading("Education", level=1)
        for item in resume.education:
            document.add_paragraph(" - ".join([value for value in [item.degree, item.institution, item.dates] if value]))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("/{resume_id}", response_model=VersionListResponse)
def list_versions(request: Request, resume_id: str) -> VersionListResponse:
    user_id = request.state.user_id
    connection = get_connection()
    try:
        rows = connection.execute(
            """
            SELECT version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                   accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after
            FROM resume_versions
            WHERE base_resume_id = ? AND user_id = ?
            ORDER BY created_at DESC
            """,
            (resume_id, user_id),
        ).fetchall()
    finally:
        connection.close()

    return VersionListResponse(
        versions=[
            VersionSummary(
                version_id=row[0],
                version_number=row[3],
                job_title=row[6] or "",
                company_name=row[5] or "",
                timestamp=row[10],
                accepted_edits_count=row[7],
                ats_score_before=float(row[11] or 0),
                ats_score_after=float(row[12] or 0),
            )
            for row in rows
        ]
    )


@router.get("/{version_id}/diff", response_model=VersionDiffResponse)
def get_version_diff(request: Request, version_id: str) -> VersionDiffResponse:
    user_id = request.state.user_id
    connection = get_connection()
    try:
        row = connection.execute(
            """
            SELECT version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                   accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after
            FROM resume_versions
            WHERE version_id = ? AND user_id = ?
            """,
            (version_id, user_id),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Version not found.")

        previous_row = connection.execute(
            """
            SELECT version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                   accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after
            FROM resume_versions
            WHERE base_resume_id = ? AND user_id = ? AND version_number < ?
            ORDER BY version_number DESC
            LIMIT 1
            """,
            (row[1], user_id, row[3]),
        ).fetchone()
    finally:
        connection.close()

    return diff_versions(row_to_version(row), row_to_version(previous_row) if previous_row else None)


@router.post("/{version_id}/export")
def export_version(
    request: Request,
    version_id: str,
    format: str = Query(..., pattern="^(pdf|docx)$"),
    template: str = Query("modern", pattern="^(modern|classic|minimal)$"),
):
    user_id = request.state.user_id
    connection = get_connection()
    try:
        row = connection.execute(
            """
            SELECT version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                   accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after,
                   preserved_docx_blob
            FROM resume_versions
            WHERE version_id = ? AND user_id = ?
            """,
            (version_id, user_id),
        ).fetchone()
        resume_row = connection.execute(
            "SELECT file_name, file_type, original_file FROM resumes WHERE id = ? AND user_id = ?",
            (row[1], user_id),
        ).fetchone() if row else None
    finally:
        connection.close()

    if not row:
        raise HTTPException(status_code=404, detail="Version not found.")

    version = row_to_version(row)

    if resume_row and resume_row[1] == "docx" and row[13]:
        if format == "docx":
            buffer = io.BytesIO(row[13])
            buffer.seek(0)
            file_stem = Path(resume_row[0] or f"resume-{version.metadata.version_number}").stem
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{file_stem}-v{version.metadata.version_number}.docx"'},
            )
        raise HTTPException(
            status_code=400,
            detail="Exact-format PDF export is not available yet. Download DOCX to preserve the original layout and fonts.",
        )

    if resume_row and resume_row[1] == "pdf":
        raise HTTPException(
            status_code=400,
            detail="Exact format-preserving export is only supported for DOCX source resumes. PDF uploads cannot be safely rewritten while keeping the original typography and layout.",
        )

    if format == "pdf":
        buffer = io.BytesIO()
        draw_pdf_resume(buffer, version.resume_json, template)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="resume-{version.metadata.version_number}.pdf"'},
        )

    buffer = build_docx_resume(version.resume_json, template)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="resume-{version.metadata.version_number}.docx"'},
    )
