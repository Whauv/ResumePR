from __future__ import annotations

import json
import io
from datetime import datetime, UTC
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, File, HTTPException, Request, UploadFile

from models.schemas import (
    ApplyEditsRequest,
    ApplyEditsResponse,
    EditSuggestion,
    GapReport,
    JobParseResponse,
    ParsedResume,
    ResumeUploadResponse,
    ResumeVersion,
    ResumeVersionListResponse,
    ResumeVersionMetadata,
)
from services.db import get_connection
from services.gap_analyzer import analyze_gap
from services.parser import parse_resume

router = APIRouter(prefix="/api/resume", tags=["resume"])


def normalize_text(value: str) -> str:
    return " ".join(value.split()).strip().lower()


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def iter_doc_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def preserve_docx_format_with_edits(file_bytes: bytes, suggestions: list[EditSuggestion]) -> bytes:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = list(iter_doc_paragraphs(document))
    paragraph_map = {normalize_text(paragraph.text): paragraph for paragraph in paragraphs if paragraph.text.strip()}

    for suggestion in suggestions:
        original = normalize_text(suggestion.original_text)
        if not original:
            continue

        paragraph = paragraph_map.get(original)
        if paragraph:
            replace_paragraph_text(paragraph, suggestion.suggested_text)
            paragraph_map[normalize_text(suggestion.suggested_text)] = paragraph
            continue

        for paragraph in paragraphs:
            text = paragraph.text or ""
            normalized_text = normalize_text(text)
            if original and original in normalized_text:
                updated_text = text.replace(suggestion.original_text, suggestion.suggested_text)
                if updated_text == text and suggestion.original_text.strip() != text.strip():
                    continue
                replace_paragraph_text(paragraph, updated_text)
                break

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def apply_suggestion_to_resume(resume_data: dict, suggestion: EditSuggestion) -> None:
    if suggestion.section == "summary":
        resume_data["summary"] = suggestion.suggested_text
        return

    if suggestion.section == "skills":
        resume_data["skills"] = [
            skill.strip()
            for skill in suggestion.suggested_text.split(",")
            if skill.strip()
        ]
        return

    if suggestion.section == "experience":
        experience_entries = resume_data.get("experience", [])
        if suggestion.section_index >= len(experience_entries):
            return
        bullets = experience_entries[suggestion.section_index].get("bullets", [])
        if 0 <= suggestion.bullet_index < len(bullets):
            bullets[suggestion.bullet_index] = suggestion.suggested_text


def build_version_metadata(
    version_number: int,
    job_id: str,
    company_name: str,
    role: str,
    accepted_count: int,
    rejected_count: int,
    created_at: str,
    ats_score_before: float = 0.0,
    ats_score_after: float = 0.0,
) -> ResumeVersionMetadata:
    return ResumeVersionMetadata(
        version_number=version_number,
        job_id=job_id,
        company_name=company_name,
        role=role,
        timestamp=created_at,
        accepted_count=accepted_count,
        rejected_count=rejected_count,
        ats_score_before=ats_score_before,
        ats_score_after=ats_score_after,
    )


@router.post("/upload", response_model=ResumeUploadResponse)
async def upload_resume(request: Request, file: UploadFile = File(...)) -> ResumeUploadResponse:
    user_id = request.state.user_id
    extension = Path(file.filename or "").suffix.lower()
    file_type_map = {".pdf": "pdf", ".docx": "docx"}
    file_type = file_type_map.get(extension)

    if not file_type:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    file_bytes = await file.read()
    try:
        parsed_data = parse_resume(file_bytes, file_type)
        parsed_resume = ParsedResume.model_validate(parsed_data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Failed to parse resume.") from exc

    resume_id = str(uuid4())
    connection = get_connection()
    try:
        connection.execute(
            "INSERT INTO resumes (id, user_id, file_name, file_type, original_file, parsed_json) VALUES (?, ?, ?, ?, ?, ?)",
            (resume_id, user_id, file.filename or "resume", file_type, file_bytes, json.dumps(parsed_resume.model_dump())),
        )
        connection.commit()
    finally:
        connection.close()

    return ResumeUploadResponse(
        resume_id=resume_id,
        file_name=file.filename or "resume",
        file_type=file_type,
        parsed_resume=parsed_resume,
    )


@router.post("/apply-edits", response_model=ApplyEditsResponse)
def apply_edits(http_request: Request, request: ApplyEditsRequest) -> ApplyEditsResponse:
    user_id = http_request.state.user_id
    if not request.accepted_edit_ids:
        raise HTTPException(status_code=400, detail="At least one accepted edit is required.")

    connection = get_connection()
    try:
        resume_row = connection.execute(
            "SELECT parsed_json, file_type, original_file FROM resumes WHERE id = ? AND user_id = ?",
            (request.resume_id, user_id),
        ).fetchone()
        if not resume_row:
            raise HTTPException(status_code=404, detail="Resume not found.")

        placeholders = ",".join("?" for _ in request.accepted_edit_ids)
        suggestion_rows = connection.execute(
            f"SELECT id, batch_id, job_id, suggestion_json FROM suggestions WHERE id IN ({placeholders}) AND resume_id = ? AND user_id = ?",
            (*request.accepted_edit_ids, request.resume_id, user_id),
        ).fetchall()
        if not suggestion_rows:
            raise HTTPException(status_code=404, detail="No matching suggestions found.")

        batch_ids = {row[1] for row in suggestion_rows}
        if len(batch_ids) != 1:
            raise HTTPException(status_code=400, detail="Accepted edits must come from a single suggestion batch.")

        batch_id = suggestion_rows[0][1]
        job_id = suggestion_rows[0][2]

        total_suggestions = connection.execute(
            "SELECT COUNT(*) FROM suggestions WHERE batch_id = ? AND user_id = ?",
            (batch_id, user_id),
        ).fetchone()[0]

        job_row = connection.execute(
            "SELECT parsed_json FROM jobs WHERE id = ? AND user_id = ?",
            (job_id, user_id),
        ).fetchone()
        if not job_row:
            raise HTTPException(status_code=404, detail="Associated job not found.")

        version_count = connection.execute(
            "SELECT COUNT(*) FROM resume_versions WHERE base_resume_id = ? AND user_id = ?",
            (request.resume_id, user_id),
        ).fetchone()[0]
    finally:
        connection.close()

    resume_data = json.loads(resume_row[0])
    source_file_type = (resume_row[1] or "").lower()
    original_file_bytes = resume_row[2]
    suggestions = [EditSuggestion.model_validate(json.loads(row[3])) for row in suggestion_rows]
    for suggestion in suggestions:
        apply_suggestion_to_resume(resume_data, suggestion)

    preserved_docx_blob = None
    if source_file_type == "docx" and original_file_bytes:
        try:
            preserved_docx_blob = preserve_docx_format_with_edits(original_file_bytes, suggestions)
        except Exception:
            preserved_docx_blob = None

    updated_resume = ParsedResume.model_validate(resume_data)
    version_id = str(uuid4())
    version_number = int(version_count) + 1
    accepted_count = len(suggestions)
    rejected_count = int(total_suggestions) - accepted_count
    job_payload = JobParseResponse.model_validate({"job_id": job_id, **json.loads(job_row[0])})
    original_resume = ParsedResume.model_validate(json.loads(resume_row[0]))
    gap_before = GapReport.model_validate(analyze_gap(original_resume.model_dump(), job_payload.model_dump()))
    gap_after = GapReport.model_validate(analyze_gap(updated_resume.model_dump(), job_payload.model_dump()))
    created_at = datetime.now(UTC).isoformat()
    metadata = build_version_metadata(
        version_number=version_number,
        job_id=job_id,
        company_name=job_payload.company_name,
        role=job_payload.job_title,
        accepted_count=accepted_count,
        rejected_count=rejected_count,
        created_at=created_at,
        ats_score_before=gap_before.overall_score,
        ats_score_after=gap_after.overall_score,
    )

    connection = get_connection()
    try:
        connection.execute(
            """
            INSERT INTO resume_versions (
                version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after
                , preserved_docx_blob
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                version_id,
                request.resume_id,
                user_id,
                version_number,
                job_id,
                job_payload.company_name,
                job_payload.job_title,
                accepted_count,
                rejected_count,
                json.dumps(updated_resume.model_dump()),
                created_at,
                gap_before.overall_score,
                gap_after.overall_score,
                preserved_docx_blob,
            ),
        )
        connection.commit()
    finally:
        connection.close()

    return ApplyEditsResponse(
        version_id=version_id,
        updated_resume=updated_resume,
        metadata=metadata,
    )


@router.get("/{resume_id}/versions", response_model=ResumeVersionListResponse)
def list_versions(request: Request, resume_id: str) -> ResumeVersionListResponse:
    user_id = request.state.user_id
    connection = get_connection()
    try:
        rows = connection.execute(
            """
            SELECT version_id, base_resume_id, user_id, version_number, job_id, company_name, role,
                   accepted_count, rejected_count, resume_json, created_at, ats_score_before, ats_score_after
            FROM resume_versions
            WHERE base_resume_id = ? AND user_id = ?
            ORDER BY version_number DESC
            """,
            (resume_id, user_id),
        ).fetchall()
    finally:
        connection.close()

    versions = [
        ResumeVersion(
            version_id=row[0],
            base_resume_id=row[1],
            metadata=build_version_metadata(
                version_number=row[3],
                job_id=row[4],
                company_name=row[5] or "",
                role=row[6] or "",
                accepted_count=row[7],
                rejected_count=row[8],
                created_at=row[10],
                ats_score_before=float(row[11] or 0),
                ats_score_after=float(row[12] or 0),
            ),
            resume_json=ParsedResume.model_validate(json.loads(row[9])),
        )
        for row in rows
    ]
    return ResumeVersionListResponse(versions=versions)
